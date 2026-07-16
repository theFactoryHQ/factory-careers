"""Generate the synthetic, non-PII resume parser fixture corpus."""

from __future__ import annotations

import io
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw
from pypdf import PdfReader, PdfWriter
import reportlab
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parent
FIXED_ZIP_TIME = (2020, 1, 1, 0, 0, 0)


def configure(pdf: canvas.Canvas) -> None:
    pdf.setAuthor("Factory Careers")
    pdf.setCreator("Factory Careers fixture generator")
    pdf.setSubject("Synthetic resume parser test fixture")
    pdf.setTitle("Factory Careers resume parser fixture")


def normalize_docx(path: Path) -> None:
    """Give generated OOXML entries stable ordering, timestamps, and compression."""
    with zipfile.ZipFile(path, "r") as source:
        entries = [(entry.filename, source.read(entry.filename)) for entry in source.infolist()]

    normalized = path.with_suffix(".normalized.docx")
    with zipfile.ZipFile(normalized, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as output:
        for name, data in sorted(entries):
            entry = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
            entry.compress_type = zipfile.ZIP_DEFLATED
            entry.external_attr = 0o600 << 16
            output.writestr(entry, data)
    normalized.replace(path)


def make_compressed_multipage() -> None:
    output = ROOT / "compressed-multipage.pdf"
    pdf = canvas.Canvas(str(output), pagesize=letter, pageCompression=1, invariant=1)
    configure(pdf)
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 720, "Factory Careers compressed multipage resume page one")
    pdf.showPage()
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 720, "Factory Careers compressed multipage resume page two")
    pdf.save()


def make_subset_font() -> None:
    font_path = Path(reportlab.__file__).resolve().parent / "fonts" / "Vera.ttf"
    if not font_path.exists():
        raise RuntimeError(f"Expected fixture font not found: {font_path}")

    pdfmetrics.registerFont(TTFont("FixtureSubset", str(font_path)))
    pdf = canvas.Canvas(
        str(ROOT / "subset-font.pdf"),
        pagesize=letter,
        pageCompression=1,
        invariant=1,
    )
    configure(pdf)
    pdf.setFont("FixtureSubset", 12)
    pdf.drawString(72, 720, "Factory Careers subset font resume")
    pdf.save()


def make_multi_column() -> None:
    pdf = canvas.Canvas(
        str(ROOT / "multi-column.pdf"),
        pagesize=letter,
        pageCompression=1,
        invariant=1,
    )
    configure(pdf)
    pdf.setFont("Helvetica", 12)
    pdf.drawString(54, 720, "Factory Careers left column")
    pdf.drawString(330, 720, "Factory Careers right column")
    pdf.drawString(54, 690, "Experience")
    pdf.drawString(330, 690, "Education")
    pdf.save()


def make_image_only() -> None:
    image = Image.new("RGB", (600, 160), "white")
    drawing = ImageDraw.Draw(image)
    drawing.rectangle((5, 5, 595, 155), outline="black", width=3)
    drawing.line((20, 50, 580, 50), fill="black", width=4)
    drawing.line((20, 100, 450, 100), fill="black", width=4)
    image_bytes = io.BytesIO()
    image.save(image_bytes, format="PNG")
    image_bytes.seek(0)

    pdf = canvas.Canvas(
        str(ROOT / "image-only.pdf"),
        pagesize=letter,
        pageCompression=1,
        invariant=1,
    )
    configure(pdf)
    pdf.drawImage(ImageReader(image_bytes), 72, 550, width=468, height=125)
    pdf.save()


def make_password_protected() -> None:
    source = io.BytesIO()
    pdf = canvas.Canvas(source, pagesize=letter, pageCompression=1, invariant=1)
    configure(pdf)
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 720, "Factory Careers password protected resume fixture")
    pdf.save()
    source.seek(0)

    reader = PdfReader(source)
    writer = PdfWriter()
    writer.append_pages_from_reader(reader)
    writer.encrypt(user_password="fixture-password", owner_password="fixture-owner-password")
    with (ROOT / "password-protected.pdf").open("wb") as output:
        writer.write(output)


def make_truncated_pdf() -> None:
    # A stable PDF header/object prefix with no xref, trailer, or EOF marker.
    (ROOT / "truncated.pdf").write_bytes(
        b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog >>\n"
    )


def make_word_documents() -> None:
    docx = Document()
    docx.add_heading("Synthetic Resume", level=1)
    docx.add_paragraph("Factory Careers DOCX resume fixture")
    docx.add_heading("Experience", level=2)
    docx.add_paragraph("Generated document parsing test content")
    docx.save(ROOT / "valid.docx")
    normalize_docx(ROOT / "valid.docx")

    legacy_source = ROOT / "legacy-source.docx"
    legacy = Document()
    legacy.add_heading("Synthetic Resume", level=1)
    legacy.add_paragraph("Factory Careers legacy DOC resume fixture")
    legacy.add_heading("Education", level=2)
    legacy.add_paragraph("Generated legacy document parsing test content")
    legacy.save(legacy_source)
    normalize_docx(legacy_source)

    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice soffice is required to generate valid.doc")

    with tempfile.TemporaryDirectory(prefix="factory-careers-soffice-") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation={profile_uri}",
                "--convert-to",
                "doc:MS Word 97",
                "--outdir",
                str(ROOT),
                str(legacy_source),
            ],
            check=True,
            capture_output=True,
            text=True,
        )

    converted = ROOT / "legacy-source.doc"
    converted.replace(ROOT / "valid.doc")
    legacy_source.unlink()


if __name__ == "__main__":
    ROOT.mkdir(parents=True, exist_ok=True)
    make_compressed_multipage()
    make_subset_font()
    make_multi_column()
    make_image_only()
    make_password_protected()
    make_truncated_pdf()
    make_word_documents()
